"""
Document Analyzer Module - Refresh v5
Extracts text from documents and uses Azure OpenAI to identify fields.

Supports: .docx, .txt
"""

import io
import json
import os
import logging
from typing import Dict, List, Optional, Any

from docx import Document
from openai import AzureOpenAI


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a .docx file using python-docx.

    Args:
        file_bytes: The .docx file as bytes

    Returns:
        Extracted text content
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))

        return '\n'.join(paragraphs)

    except Exception as e:
        logging.error(f"Error extracting text from docx: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('latin-1')
    except Exception as e:
        logging.error(f"Error extracting text from txt: {e}")
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from a document based on file extension.

    Args:
        file_bytes: The file as bytes
        filename: Original filename (used to detect type)

    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()

    if filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith('.txt'):
        return extract_text_from_txt(file_bytes)
    elif filename_lower.endswith('.doc'):
        raise ValueError("Legacy .doc format not supported. Please convert to .docx")
    else:
        # Try as plain text
        return extract_text_from_txt(file_bytes)


# System prompt for document analysis
ANALYSIS_SYSTEM_PROMPT = """You are a document analysis assistant for a school district. Analyze business documents and extract ATOMIC, EDITABLE fields — the specific values a user would want to change when refreshing this document for a new school year or new recipient.

CRITICAL RULES FOR FIELD EXTRACTION:
1. Extract ATOMIC values, NOT entire sentences or paragraphs.
   - WRONG: current_value = "The training will begin on Monday, January 2 and end on Thursday, January 5"
   - RIGHT: Two separate fields:
     - field_name: "start_date", current_value: "Monday, January 2"
     - field_name: "end_date", current_value: "Thursday, January 5"

2. Every field must be a SINGLE, STANDALONE value that a user can edit independently.
   - Dates → extract each date separately (start_date, end_date, deadline, event_date, etc.)
   - Names → extract each name separately (principal_name, superintendent_name, recipient_name)
   - Titles → extract each title separately (principal_title, sender_title)
   - Locations → extract each location (school_name, venue, address)
   - Times → extract each time (start_time, end_time)
   - School year → extract as its own field (school_year: "2025-2026")

3. The "body" field should contain ONLY the main prose content, with specific values already extracted as separate fields. Use {{field_name}} placeholders in the body where extracted values appear.
   - Example body: "The training will begin on {{start_date}} and end on {{end_date}} at {{location}}."

4. Do NOT extract boilerplate text that never changes (e.g., "Sincerely," or "MEMORANDUM").

5. Do NOT return entire paragraphs as a single field. Break them into their component editable values.

When given a document, you must:
1. Identify the document type (memo, letter, resolution, report, general correspondence)
2. Extract ALL specific, editable values as individual atomic fields
3. Return structured JSON

For each field, provide:
- field_name: A standardized snake_case name (e.g., "recipient_name", "start_date", "school_year")
- field_label: A human-readable label (e.g., "Recipient Name", "Start Date", "School Year")
- current_value: The EXACT value as it appears in the document (just the value, not the surrounding sentence)
- field_type: One of "text", "date", "multiline", "list"
- required: true/false

Common document types and their typical ATOMIC fields:

MEMO:
- recipient_name (TO: person/group name)
- recipient_title (TO: person's title, if present)
- sender_name (FROM: person name)
- sender_title (FROM: person's title, if present)
- subject (RE: or SUBJECT: line)
- date (DATE: value)
- body (main content with {{placeholders}} for extracted values - multiline)
- cc (CC: names - optional)

FORMAL LETTER:
- date (letter date)
- recipient_name
- recipient_title
- recipient_organization
- recipient_address
- salutation (Dear...)
- body (multiline, with {{placeholders}})
- closing (Sincerely, etc.)
- sender_name
- sender_title

SCHOOL DISTRICT DOCUMENTS (common additional fields):
- school_year (e.g., "2025-2026")
- principal_name
- superintendent_name
- school_name
- event_date / start_date / end_date
- deadline
- grade_levels
- contact_name / contact_email / contact_phone

Return ONLY valid JSON:
{
  "document_type": "memo|letter|resolution|report|correspondence|other",
  "document_type_display": "Human readable type name",
  "confidence": 0.0-1.0,
  "fields": [
    {
      "field_name": "string",
      "field_label": "string",
      "current_value": "string",
      "field_type": "text|date|multiline|list",
      "required": true|false
    }
  ],
  "summary": "Brief one-sentence description of the document"
}

EXAMPLE — Given a memo containing:
"TO: Board of Education
FROM: Dr. Jane Smith, Superintendent
RE: Professional Development Training
DATE: January 15, 2025
The district will hold mandatory professional development training beginning Monday, January 20, 2025 and ending Friday, January 24, 2025 at Port Jefferson Middle School."

You should extract:
- recipient_name: "Board of Education"
- sender_name: "Dr. Jane Smith"
- sender_title: "Superintendent"
- subject: "Professional Development Training"
- date: "January 15, 2025"
- event_name: "professional development training"
- start_date: "Monday, January 20, 2025"
- end_date: "Friday, January 24, 2025"
- location: "Port Jefferson Middle School"
- body: (the full paragraph text with {{placeholders}})

NOT a single field with the entire paragraph as its value."""


def analyze_document_with_llm(
    text: str,
    azure_endpoint: str,
    azure_api_key: str,
    azure_deployment: str,
    api_version: str
) -> Dict[str, Any]:
    """
    Use Azure OpenAI to analyze document and extract fields.

    Args:
        text: The document text
        azure_endpoint: Azure OpenAI endpoint URL
        azure_api_key: Azure OpenAI API key
        azure_deployment: Deployment name
        api_version: API version

    Returns:
        Dictionary with document analysis results
    """
    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=azure_api_key,
        api_version=api_version
    )

    # Truncate if too long (increased for better field extraction)
    max_chars = 12000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Document truncated for analysis...]"

    try:
        response = client.chat.completions.create(
            model=azure_deployment,
            messages=[
                {"role": "system", "content": ANALYSIS_SYSTEM_PROMPT},
                {"role": "user", "content": f"Analyze this document and extract its variable fields:\n\n{text}"}
            ],
            temperature=0.1,
            max_tokens=2000,
            response_format={"type": "json_object"}
        )

        result_text = response.choices[0].message.content
        return json.loads(result_text)

    except Exception as e:
        logging.error(f"Error calling Azure OpenAI: {e}")
        raise


def analyze_document(
    file_bytes: bytes,
    filename: str,
    azure_endpoint: Optional[str] = None,
    azure_api_key: Optional[str] = None,
    azure_deployment: Optional[str] = None,
    api_version: Optional[str] = None
) -> Dict[str, Any]:
    """
    Full document analysis pipeline.

    Args:
        file_bytes: The document file as bytes
        filename: Original filename
        azure_endpoint: Azure OpenAI endpoint (or env var)
        azure_api_key: Azure OpenAI API key (or env var)
        azure_deployment: Deployment name (or env var)
        api_version: API version (or env var)

    Returns:
        Analysis results with:
        - document_type, document_type_display, confidence
        - fields: Array of extracted fields
        - summary: Brief description
    """
    # Get Azure OpenAI config
    endpoint = azure_endpoint or os.environ.get('AZURE_OPENAI_ENDPOINT')
    api_key = azure_api_key or os.environ.get('AZURE_OPENAI_KEY')
    deployment = azure_deployment or os.environ.get('AZURE_OPENAI_DEPLOYMENT', 'gpt-4o-mini')
    version = api_version or os.environ.get('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')

    if not endpoint or not api_key:
        raise ValueError(
            "Azure OpenAI credentials not configured. "
            "Set AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY environment variables."
        )

    # Extract text
    text = extract_text(file_bytes, filename)

    if not text.strip():
        raise ValueError("Could not extract text from document. File may be empty or corrupted.")

    # Analyze with LLM
    analysis = analyze_document_with_llm(text, endpoint, api_key, deployment, version)

    # Add metadata
    analysis['filename'] = filename
    analysis['original_text_preview'] = text[:500] + "..." if len(text) > 500 else text

    return analysis


def generate_results_card(analysis: Dict[str, Any]) -> Dict:
    """
    Generate an Adaptive Card showing analysis results.

    Args:
        analysis: Document analysis results

    Returns:
        Adaptive Card JSON
    """
    facts = []
    for field in analysis.get('fields', [])[:10]:  # Limit to 10 fields
        value = field.get('current_value', '')
        if len(value) > 50:
            value = value[:50] + "..."
        facts.append({
            "title": field.get('field_label', field.get('field_name', 'Field')),
            "value": value
        })

    return {
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "type": "AdaptiveCard",
        "version": "1.5",
        "body": [
            {
                "type": "TextBlock",
                "text": "Document Analyzed",
                "weight": "Bolder",
                "size": "Large",
                "color": "Good",
                "wrap": True
            },
            {
                "type": "FactSet",
                "facts": [
                    {"title": "Document Type:", "value": analysis.get('document_type_display', 'Unknown')},
                    {"title": "Confidence:", "value": f"{int(analysis.get('confidence', 0) * 100)}%"}
                ]
            },
            {
                "type": "TextBlock",
                "text": analysis.get('summary', ''),
                "wrap": True,
                "isSubtle": True,
                "spacing": "Small"
            },
            {
                "type": "TextBlock",
                "text": "Extracted Fields",
                "weight": "Bolder",
                "spacing": "Medium"
            },
            {
                "type": "FactSet",
                "facts": facts
            }
        ],
        "actions": [
            {
                "type": "Action.Submit",
                "title": "Update These Fields",
                "style": "positive",
                "data": {"action": "update_fields"}
            },
            {
                "type": "Action.Submit",
                "title": "Generate with Same Values",
                "data": {"action": "generate_same"}
            }
        ]
    }


def generate_input_card(analysis: Dict[str, Any]) -> Dict:
    """
    Generate an Adaptive Card with input fields for updating.

    Args:
        analysis: Document analysis results

    Returns:
        Adaptive Card JSON
    """
    body = [
        {
            "type": "TextBlock",
            "text": "Enter New Values",
            "weight": "Bolder",
            "size": "Large",
            "wrap": True
        },
        {
            "type": "TextBlock",
            "text": "Update the fields you want to change.",
            "wrap": True,
            "isSubtle": True
        }
    ]

    for field in analysis.get('fields', []):
        field_name = field.get('field_name', 'field')
        field_label = field.get('field_label', field_name)
        field_type = field.get('field_type', 'text')
        current_value = field.get('current_value', '')

        body.append({
            "type": "TextBlock",
            "text": field_label,
            "weight": "Bolder",
            "spacing": "Medium"
        })

        if field_type == 'date':
            body.append({
                "type": "Input.Date",
                "id": field_name,
                "value": current_value
            })
        elif field_type == 'multiline':
            placeholder = current_value[:50] + "..." if len(current_value) > 50 else current_value
            body.append({
                "type": "Input.Text",
                "id": field_name,
                "isMultiline": True,
                "value": current_value,
                "placeholder": f"Current: {placeholder}"
            })
        else:
            body.append({
                "type": "Input.Text",
                "id": field_name,
                "value": current_value,
                "placeholder": f"Current: {current_value}"
            })

    return {
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "type": "AdaptiveCard",
        "version": "1.5",
        "body": body,
        "actions": [
            {
                "type": "Action.Submit",
                "title": "Generate Document",
                "style": "positive",
                "data": {"action": "generate"}
            },
            {
                "type": "Action.Submit",
                "title": "Cancel",
                "data": {"action": "cancel"}
            }
        ]
    }
