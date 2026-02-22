"""
Document Generator Module - Refresh v6
Clean document generation using python-docx.

v6 changes:
- New function: generate_from_synthesis() — generates from comparative analysis
- All v5 generation functions preserved for backward compatibility / fallback

Supports:
- Memos
- Formal Letters
- Synthesis-based generation (new for v6)
- Generic documents with field replacement
"""

import io
import os
import re
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from openai import AzureOpenAI


class DocumentGenerator:
    """
    Unified document generator using python-docx.

    Usage:
        generator = DocumentGenerator()
        doc_bytes = generator.generate_memo(fields)

        # Or with a template:
        generator = DocumentGenerator(template_bytes)
        doc_bytes = generator.replace_placeholders(fields)
    """

    def __init__(self, template_bytes: Optional[bytes] = None):
        """
        Initialize the document generator.

        Args:
            template_bytes: Optional .docx template file as bytes.
                           If provided, will use as base document.
        """
        if template_bytes:
            self.doc = Document(io.BytesIO(template_bytes))
        else:
            self.doc = Document()
            self._setup_default_styles()

    def _setup_default_styles(self):
        """Set up default document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Century Schoolbook'
        font.size = Pt(11)

        for section in self.doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def generate_memo(self, fields: Dict[str, Any]) -> bytes:
        """
        Generate a memorandum document.

        Args:
            fields: Dictionary containing:
                - recipient (str): TO field
                - sender (str): FROM field
                - subject (str): RE/Subject field
                - date (str): Date field
                - body (str or List[str]): Body paragraphs
                - cc (str, optional): CC field

        Returns:
            Generated .docx file as bytes
        """
        for _ in range(3):
            self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("MEMORANDUM")
        run.bold = True
        run.underline = True
        run.font.size = Pt(12)
        run.font.name = 'Century Schoolbook'

        self.doc.add_paragraph()

        self._add_memo_field("TO:", fields.get('recipient', ''))
        self.doc.add_paragraph()
        self._add_memo_field("FROM:", fields.get('sender', ''))
        self.doc.add_paragraph()
        self._add_memo_field("RE:", fields.get('subject', ''))
        self.doc.add_paragraph()
        self._add_memo_field("DATE:", fields.get('date', datetime.now().strftime('%B %d, %Y').upper()))

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        body = fields.get('body', [])
        if isinstance(body, str):
            body = [p.strip() for p in body.split('\n\n') if p.strip()]

        for para_text in body:
            para = self.doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = 'Century Schoolbook'
            run.font.size = Pt(11)
            self.doc.add_paragraph()

        cc = fields.get('cc')
        if cc:
            for _ in range(3):
                self.doc.add_paragraph()
            para = self.doc.add_paragraph()
            run = para.add_run(f"CC: {cc}")
            run.font.name = 'Century Schoolbook'
            run.font.size = Pt(10)

        return self._to_bytes()

    def _add_memo_field(self, label: str, value: str):
        """Add a memo header field (TO:, FROM:, etc.)."""
        para = self.doc.add_paragraph()
        label_run = para.add_run(label)
        label_run.font.name = 'Century Schoolbook'
        label_run.font.size = Pt(11)
        para.add_run("\t")
        value_run = para.add_run(value)
        value_run.font.name = 'Century Schoolbook'
        value_run.font.size = Pt(11)

    def generate_letter(self, fields: Dict[str, Any]) -> bytes:
        """
        Generate a formal business letter.

        Args:
            fields: Dictionary containing:
                - date, recipient_name, recipient_title, recipient_organization,
                  recipient_address, salutation, body, closing, sender_name, sender_title

        Returns:
            Generated .docx file as bytes
        """
        for _ in range(4):
            self.doc.add_paragraph()

        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(fields.get('date', datetime.now().strftime('%B %d, %Y')))
        date_run.font.name = 'Century Schoolbook'
        date_run.font.size = Pt(11)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        if fields.get('recipient_name'):
            self._add_line(fields['recipient_name'])
        if fields.get('recipient_title'):
            self._add_line(fields['recipient_title'])
        if fields.get('recipient_organization'):
            self._add_line(fields['recipient_organization'])
        if fields.get('recipient_address'):
            for line in fields['recipient_address'].split('\n'):
                if line.strip():
                    self._add_line(line.strip())

        self.doc.add_paragraph()

        salutation = fields.get('salutation', 'Dear Sir or Madam:')
        self._add_line(salutation)
        self.doc.add_paragraph()

        body = fields.get('body', [])
        if isinstance(body, str):
            body = [p.strip() for p in body.split('\n\n') if p.strip()]

        for para_text in body:
            para = self.doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = 'Century Schoolbook'
            run.font.size = Pt(11)
            self.doc.add_paragraph()

        closing = fields.get('closing', 'Sincerely,')
        self._add_line(closing)

        for _ in range(3):
            self.doc.add_paragraph()

        if fields.get('sender_name'):
            self._add_line(fields['sender_name'])
        if fields.get('sender_title'):
            self._add_line(fields['sender_title'])

        return self._to_bytes()

    def _add_line(self, text: str):
        """Add a single line of text."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Century Schoolbook'
        run.font.size = Pt(11)

    def replace_placeholders(self, fields: Dict[str, Any]) -> bytes:
        """
        Replace {{placeholder}} tokens in the document with actual values.

        Args:
            fields: Dictionary of field_name -> value

        Returns:
            Modified .docx file as bytes
        """
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, fields)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, fields)

        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, fields)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, fields)

        return self._to_bytes()

    def _replace_in_paragraph(self, para, fields: Dict[str, Any]):
        """Replace placeholders in a single paragraph."""
        full_text = para.text

        if '{{' not in full_text:
            return

        for field_name, value in fields.items():
            placeholder = f"{{{{{field_name}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value) if value else '')

        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            for run in para.runs:
                run.text = ''

            para.runs[0].text = full_text
            para.runs[0].font.name = font_name
            para.runs[0].font.size = font_size
            para.runs[0].bold = bold
            para.runs[0].italic = italic

    def generate_from_analysis(self, document_type: str, fields: Dict[str, Any]) -> bytes:
        """
        Generate a document based on analyzed document type.

        Args:
            document_type: Type of document (memo, letter, etc.)
            fields: Field values

        Returns:
            Generated .docx file as bytes
        """
        doc_type = document_type.lower()

        if doc_type == 'memo' or doc_type == 'memorandum':
            return self.generate_memo(fields)
        elif doc_type == 'letter' or doc_type == 'formal_letter':
            return self.generate_letter(fields)
        else:
            return self.generate_letter(fields)

    def _to_bytes(self) -> bytes:
        """Save document to bytes."""
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def generate_document(
    document_type: str,
    fields: Dict[str, Any],
    template_bytes: Optional[bytes] = None
) -> bytes:
    """
    Convenience function to generate a document.

    Args:
        document_type: Type of document (memo, letter)
        fields: Field values
        template_bytes: Optional template to use

    Returns:
        Generated .docx file as bytes
    """
    generator = DocumentGenerator(template_bytes)

    if template_bytes:
        return generator.replace_placeholders(fields)
    else:
        return generator.generate_from_analysis(document_type, fields)


def generate_filename(document_type: str, fields: Dict[str, Any]) -> str:
    """
    Generate a filename for the document.

    Args:
        document_type: Type of document
        fields: Field values (may contain subject or other identifying info)

    Returns:
        Generated filename
    """
    subject = fields.get('subject', fields.get('re', document_type.title()))
    safe_subject = "".join(c for c in str(subject)[:30] if c.isalnum() or c in ' -_')
    date_str = datetime.now().strftime('%m%d%Y')
    return f"{document_type.title()} - {safe_subject} - {date_str}.docx"


# =============================================================================
# v6: SYNTHESIS-BASED GENERATION (NEW)
# =============================================================================

SYNTHESIS_GENERATION_PROMPT = """You are a document generator for a school district.
You have a comparative analysis of how this document has evolved over time,
the full text of the most recent version, and the user's specific requests.

Your task: Generate the COMPLETE TEXT of a new version of this document.

Rules:
1. Use the most recent version as the base structure and tone.
2. Apply all predicted variable changes from the analysis.
3. Apply any user-requested changes.
4. If organizational context is provided, proactively incorporate relevant
   organizational changes into the document. For example, if a new hire
   announcement mentions a new assistant principal, update the document
   to reference that person. If a budget memo mentions a new program,
   add it where appropriate. Cite the source in changes_applied.
5. Preserve emerging elements (recently added sections).
6. For unpredictable values you cannot determine, use [PLACEHOLDER] markers
   and include them in the flags array.
7. Maintain the document's original tone and style.
8. Target year/period: {target_info}

Return ONLY valid JSON:
{{
  "generated_text": "The complete document text...",
  "changes_applied": ["Changed X to Y", "Updated Z"],
  "flags": [
    {{"field": "field_name", "reason": "why this needs input", "placeholder": "[TEXT]"}}
  ],
  "suggested_filename": "Document Name - Year.docx"
}}"""


def generate_from_synthesis(
    family_analysis: Dict[str, Any],
    base_document_text: str,
    user_changes: str = '',
    target_year: str = '',
    organizational_context: str = '',
    azure_endpoint: Optional[str] = None,
    azure_api_key: Optional[str] = None,
    azure_deployment: Optional[str] = None,
    api_version: Optional[str] = None
) -> Dict[str, Any]:
    """
    Generate a new document grounded in comparative family analysis,
    optionally incorporating organizational context.

    Uses the most recent version as a base, applies predicted variable
    updates from the analysis, incorporates organizational context
    (discovered changes in the organization), and applies user-requested changes.

    Args:
        family_analysis: The output from analyze_document_family()
        base_document_text: Full plain text of the most recent version
        user_changes: Natural language description of user-requested changes
        target_year: Target year/period (e.g., "2026-2027")
        organizational_context: Summary of relevant organizational changes
            discovered from context documents (e.g., "New AP Dr. Johnson hired;
            1:1 device program expanding"). Empty string if none found.
        azure_endpoint, etc.: Azure OpenAI config (or from env vars)

    Returns:
        Dictionary with:
        - generated_text: str (the full document text)
        - changes_applied: List[str] (what was changed)
        - flags: List[dict] (items needing user input)
        - suggested_filename: str
    """
    endpoint = azure_endpoint or os.environ.get('AZURE_OPENAI_ENDPOINT')
    api_key = azure_api_key or os.environ.get('AZURE_OPENAI_KEY')
    deployment = azure_deployment or os.environ.get(
        'AZURE_OPENAI_DEPLOYMENT_LARGE',
        os.environ.get('AZURE_OPENAI_DEPLOYMENT', 'gpt-4o-mini')
    )
    version = api_version or os.environ.get('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')

    if not endpoint or not api_key:
        raise ValueError(
            "Azure OpenAI credentials not configured. "
            "Set AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY environment variables."
        )

    target_info = target_year or "next iteration"
    system_prompt = SYNTHESIS_GENERATION_PROMPT.format(target_info=target_info)

    # Build user message with all context
    user_msg = f"COMPARATIVE ANALYSIS:\n{json.dumps(family_analysis, indent=2)}\n\n"
    user_msg += f"MOST RECENT VERSION TEXT:\n{base_document_text[:8000]}\n\n"
    if organizational_context:
        user_msg += f"ORGANIZATIONAL CONTEXT (recently discovered changes in the organization — incorporate where relevant):\n{organizational_context}\n\n"
    if user_changes:
        user_msg += f"USER REQUESTED CHANGES:\n{user_changes}\n\n"
    user_msg += "Generate the complete new version."

    client = AzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=version
    )

    try:
        response = client.chat.completions.create(
            model=deployment,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.3,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )

        result = json.loads(response.choices[0].message.content)

        logging.info(
            f"Synthesis generation complete: "
            f"{len(result.get('changes_applied', []))} changes, "
            f"{len(result.get('flags', []))} flags"
        )

        return result

    except Exception as e:
        logging.error(f"Error generating from synthesis: {e}")
        raise
